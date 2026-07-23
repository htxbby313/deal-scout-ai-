from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document


DOCX = "Acquisition CRM MVP - Developer Build Spec.docx"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def attr(el, name):
    return el.attrib.get(f"{{{NS['w']}}}{name}")


def main():
    doc = Document(DOCX)
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"sections={len(doc.sections)}")
    print(f"headings={[p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]}")

    with ZipFile(DOCX) as zf:
        xml = zf.read("word/document.xml")
    root = ET.fromstring(xml)

    tbls = root.findall(".//w:tbl", NS)
    print(f"xml_tables={len(tbls)}")
    for idx, tbl in enumerate(tbls, start=1):
        tbl_w = tbl.find(".//w:tblPr/w:tblW", NS)
        ind = tbl.find(".//w:tblPr/w:tblInd", NS)
        grid_cols = [attr(c, "w") for c in tbl.findall(".//w:tblGrid/w:gridCol", NS)]
        row_count = len(tbl.findall("./w:tr", NS))
        print(
            f"table_{idx}: rows={row_count} tblW={attr(tbl_w, 'w') if tbl_w is not None else None} "
            f"tblInd={attr(ind, 'w') if ind is not None else None} grid={','.join(grid_cols)}"
        )


if __name__ == "__main__":
    main()
