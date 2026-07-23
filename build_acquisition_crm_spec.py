from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Acquisition CRM MVP - Developer Build Spec.docx"


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Acquisition CRM MVP Build Spec")
    set_font(run, size=9, color="6B7280")


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Acquisition CRM MVP")
    set_font(r, size=24, color="0B2545", bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run("Developer build spec for a real estate acquisition pipeline, automations, and lead follow-up system")
    set_font(r2, size=11, color="475569", italic=True)

    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [3120, 3120, 3120])
    labels = [
        ("Primary outcome", "No seller lead falls through the cracks"),
        ("MVP posture", "Usable CRM first, AI enhancements second"),
        ("Core workflow", "Lead intake -> contact -> intake -> evaluation -> offer -> contract/follow-up"),
    ]
    for idx, (label, value) in enumerate(labels):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "F4F6F9")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label.upper())
        set_font(r, size=8, color="64748B", bold=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=10.5, color="0F172A", bold=True)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r)


def add_table(doc, headers, rows, widths, header_fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=9.5, color="0B2545", bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            set_font(r, size=9.5)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_section_overview(doc):
    doc.add_heading("1. Product Goal", level=1)
    doc.add_paragraph(
        "Build a CRM that receives seller leads, verifies them, moves them through a repeatable acquisition workflow, "
        "creates the right follow-up tasks automatically, and gives the acquisition team a clear daily work queue."
    )
    doc.add_heading("MVP Success Criteria", level=2)
    add_bullets(doc, [
        "Every new seller lead has a source, required contact details, property address, owner assignment, and current status.",
        "Duplicate phone numbers and duplicate property addresses are detected before a new record is created.",
        "Securing a lead triggers welcome outreach, activity logging, and placement into Day 1 of the contact board.",
        "Contact attempts, seller responses, intake completion, deal evaluation, offers, negotiation, contracts, and follow-up are tracked as state changes.",
        "Automations are event-driven and stop or change when the seller responds, an offer is accepted, or a deal is marked dead.",
    ])


def add_domain_model(doc):
    doc.add_heading("2. Core Data Model", level=1)
    add_table(doc, ["Entity", "Purpose", "Key Fields"], [
        ("Lead", "Seller and property record moving through the CRM.", "name, phone, email, property_address, source, status, owner_id, created_at, last_activity_at"),
        ("Property Intake", "Structured facts needed to evaluate the property.", "bedrooms, bathrooms, condition, repairs_needed, occupancy, motivation, timeline, mortgage, notes"),
        ("Activity", "Chronological history of system and user actions.", "lead_id, type, direction, channel, body, outcome, created_by, created_at"),
        ("Task", "Work item for calls, follow-ups, review, and overdue recovery.", "lead_id, assignee_id, task_type, due_at, status, completed_at, result"),
        ("Automation Run", "Audit trail for triggered workflows.", "lead_id, trigger, step, state, scheduled_for, executed_at, canceled_reason"),
        ("Buy Box", "Admin-configured investment criteria.", "market, zip_codes, max_price, min_equity, max_rehab, cash_flow, cap_rate, arv_rule"),
        ("Deal Evaluation", "Computed acquisition recommendation.", "lead_id, strategy, deal_score, arv, rehab_estimate, max_offer, rationale"),
        ("Offer", "Offer lifecycle and seller negotiation.", "lead_id, amount, status, sent_at, accepted_at, rejected_at, counter_amount, final_agreed_price"),
        ("Contract", "Generated agreement package and signature state.", "lead_id, offer_id, contract_type, document_url, signature_status, signed_at"),
    ], [1400, 2500, 5460])

    doc.add_heading("Required Lead Fields", level=2)
    add_bullets(doc, ["Name", "Phone", "Email", "Property Address", "Lead Source", "Date Created"])


def add_states(doc):
    doc.add_heading("3. Lead Statuses and State Rules", level=1)
    add_table(doc, ["Status", "Entry Trigger", "Exit Condition"], [
        ("New Lead", "Lead received from PPC, Facebook, website, manual entry, CSV, or API.", "Verification starts."),
        ("Verified", "Phone, email, property duplicate, and phone duplicate checks complete.", "User clicks Secure Lead or auto-secure rule runs."),
        ("Secured", "Welcome SMS/email sent and activity history created.", "Lead moves to Day 1."),
        ("Day 1-5", "Lead is in the 5-day contact sequence.", "Seller responds or sequence completes."),
        ("Contacted", "Seller responds.", "Intake requested."),
        ("Intake Requested", "Acquisition manager asks seller to complete intake.", "Intake submitted or follow-up continues."),
        ("Intake Received", "Property intake form submitted.", "Evaluation starts."),
        ("Evaluating", "System analyzes property against buy boxes.", "Strategy and score generated."),
        ("Offer Sent", "Offer has been delivered to seller.", "Negotiating, accepted, rejected, or no response follow-up."),
        ("Negotiating", "Counter offer or active price discussion exists.", "Accepted, rejected, or dead deal."),
        ("Contract Signed", "Purchase agreement or assignment is signed.", "Closed or contract management."),
        ("Long Follow-Up", "No seller response after contact sequence.", "Seller responds or 12 months pass."),
        ("Forever Follow-Up", "Month 13 begins after long follow-up.", "Seller responds; never auto-delete."),
        ("Dead Deal", "Rejected, disqualified, bad data, or admin closes out.", "Manual reopen only."),
        ("Closed", "Deal completed.", "Final reporting only."),
    ], [1550, 4100, 3710])


def add_workflows(doc):
    doc.add_heading("4. Workflow and Automation Logic", level=1)

    doc.add_heading("Lead Intake", level=2)
    add_numbered(doc, [
        "Receive lead from PPC, Facebook, website form, manual entry, CSV import, or API.",
        "Normalize contact fields and property address.",
        "Validate phone and email format.",
        "Check duplicate phone number and duplicate property address.",
        "Merge with existing record when duplicate confidence is high; otherwise create a new lead.",
    ])

    doc.add_heading("Secure Lead Event", level=2)
    add_table(doc, ["Trigger", "System Actions"], [
        ("Secure Lead button clicked", "Send welcome SMS, send welcome email, create activity entries, create Day 1 call task, move lead into Day 1 Contact Board."),
        ("Seller responds", "Cancel open contact-sequence automations, assign acquisition manager, create response activity, move to Contacted."),
        ("Intake form submitted", "Store intake, create activity, generate property analysis, score deal, recommend strategy, notify acquisition manager."),
        ("Offer sent", "Create offer activity, start 4-day offer follow-up sequence unless seller responds."),
        ("Offer accepted", "Stop offer follow-up, notify manager/admin, create contract-generation task, move to Contract Stage."),
    ], [2500, 6860])

    doc.add_heading("5-Day Contact Sequence", level=2)
    add_table(doc, ["Day", "Tasks", "Automation Rule"], [
        ("Day 1", "Call seller, leave voicemail if no answer, send text, send email.", "Create tasks immediately after Secure Lead."),
        ("Day 2", "Call seller, send text, send email.", "Create if no seller response after Day 1."),
        ("Day 3", "Call seller, send text, send email.", "Create if no seller response after Day 2."),
        ("Day 4", "Call seller, send text, send email.", "Create if no seller response after Day 3."),
        ("Day 5", "Call seller, send text, send email.", "Create if no seller response after Day 4; then move to Long Follow-Up."),
    ], [900, 4860, 3600])

    doc.add_heading("Offer Follow-Up", level=2)
    doc.add_paragraph(
        "When an offer is sent and the seller does not respond, run a four-day follow-up sequence. Each day creates a call task, sends a text, "
        "and sends an email. The sequence stops immediately when the seller replies, rejects, accepts, or the lead is manually paused."
    )

    doc.add_heading("Long-Term and Forever Follow-Up", level=2)
    add_bullets(doc, [
        "If a seller never responds after Day 5, move the lead into 12 Month Follow-Up.",
        "Every month, automatically create a phone call task and send a text/email touch.",
        "On Month 13, move the lead into Forever Follow-Up.",
        "Forever Follow-Up sends SMS and email every 30 days until the seller responds or the lead is manually suppressed.",
        "Never delete a lead automatically; preserve history for future acquisition opportunities.",
    ])


def add_buy_box(doc):
    doc.add_heading("5. Buy Box and Deal Scoring Engine", level=1)
    doc.add_paragraph(
        "Admins need configurable buy boxes so the CRM can score a property consistently before the acquisition manager spends time drafting an offer."
    )
    add_table(doc, ["Input", "Example Rule", "Scoring Use"], [
        ("Maximum Purchase Price", "Purchase price <= admin threshold.", "Reject or cap max offer if exceeded."),
        ("Minimum Equity", "Estimated equity >= required percentage or dollar amount.", "Higher equity increases score."),
        ("Maximum Rehab Cost", "Rehab estimate <= market-specific max.", "High rehab lowers score or changes strategy."),
        ("Cash Flow", "Projected monthly cash flow >= threshold.", "Rental strategy qualifier."),
        ("Cap Rate", "Projected cap rate >= threshold.", "Rental/investor strategy qualifier."),
        ("ARV", "After-repair value supports target spread.", "Flip/wholesale qualifier."),
        ("Market and Zip Code", "Property is inside active markets.", "Outside-market records may become referrals or rejects."),
    ], [2100, 3400, 3860])

    doc.add_heading("Recommended Scoring Output", level=2)
    add_bullets(doc, [
        "Deal Score: 0-100.",
        "Recommended Strategy: Flip, Rental, Wholesale, Creative Finance, Agent Referral, or Reject.",
        "Rationale: short explanation of why the strategy was selected.",
        "Missing Data: required fields needed before the score can be trusted.",
        "Next Action: draft offer, request more seller info, assign to agent referral, or reject.",
    ])


def add_screens_permissions(doc):
    doc.add_heading("6. Application Screens", level=1)
    add_table(doc, ["Screen", "Primary Jobs"], [
        ("Lead Inbox", "View new/verified leads, run duplicate review, secure lead, assign owner."),
        ("Kanban Pipeline", "Drag leads through statuses, see overdue tasks, open lead detail quickly."),
        ("Lead Detail", "Contact info, property data, activity timeline, notes, tasks, files, offers, contracts."),
        ("Property Intake Form", "Seller-facing or internal form for bedrooms, bathrooms, condition, repairs, occupancy, motivation, timeline, mortgage, and notes."),
        ("Evaluation Workspace", "Buy box fit, deal score, strategy recommendation, max offer, missing data."),
        ("Offer Workspace", "Draft offer, send offer, track negotiation notes, counter offers, final agreed price."),
        ("Contract Workspace", "Generate purchase agreement, assignment, disclosures, and signature packet."),
        ("Dashboard", "New leads, today's calls, offers sent, contracts pending, closed deals, response time, conversion rate, dead leads, follow-ups due today."),
        ("Admin Settings", "Users, roles, lead sources, buy boxes, templates, automation rules, integrations."),
    ], [1900, 7460])

    doc.add_heading("Permissions", level=2)
    add_table(doc, ["Role", "Allowed", "Restricted"], [
        ("Admin", "Everything: users, settings, automations, financial rules, reporting, delete/archive controls.", "None except explicit owner-only controls if desired."),
        ("Acquisition Manager", "Manage leads, create offers, view analytics, evaluate deals, negotiate, request contracts.", "Cannot change global automations or admin buy box rules unless granted."),
        ("Virtual Assistant", "Call leads, update notes, complete tasks, send approved follow-ups, update non-financial fields.", "No financial permissions, no buy box edits, no offer amount approval, no contract generation."),
    ], [1600, 5060, 2700])


def add_notifications_ai_build(doc):
    doc.add_heading("7. Notifications", level=1)
    add_bullets(doc, [
        "New lead arrives.",
        "Seller replies.",
        "Property intake is completed.",
        "Deal evaluation is ready.",
        "Offer is accepted.",
        "Contract is signed.",
        "Follow-up task is overdue.",
    ])

    doc.add_heading("8. Optional AI Features", level=1)
    add_table(doc, ["Feature", "MVP Behavior", "Guardrail"], [
        ("Conversation Summary", "Summarize seller calls, texts, and emails into a timeline note.", "Do not overwrite original activity records."),
        ("Motivation Analysis", "Classify motivation as high, medium, or low with reasoning.", "Show confidence and missing context."),
        ("Repair Estimate Assist", "Estimate repair category from notes or photos.", "Require human review before using in offer math."),
        ("Close Likelihood", "Predict likelihood of closing based on response speed and engagement.", "Do not hide low-score leads automatically."),
        ("Follow-Up Drafting", "Draft personalized SMS/email follow-ups.", "Require approval for new templates and sensitive language."),
        ("Strategy Recommendation", "Recommend flip, rental, wholesale, creative finance, referral, or reject.", "Explain scoring inputs and missing data."),
    ], [2100, 4300, 2960])

    doc.add_heading("9. Suggested Build Order", level=1)
    add_numbered(doc, [
        "User authentication and roles.",
        "Lead database and CRM lead detail page.",
        "Lead intake sources, validation, and duplicate detection.",
        "Kanban pipeline with drag-and-drop statuses.",
        "Activity history and task scheduler for calls and follow-ups.",
        "Automated SMS and email workflow engine.",
        "Property intake form and intake submission event.",
        "Deal evaluation and configurable buy box engine.",
        "Offer generation, offer status tracking, and negotiation notes.",
        "Long-term and forever follow-up automations.",
        "Dashboard and analytics.",
        "AI-powered lead scoring, summaries, and strategy recommendations.",
    ])

    doc.add_heading("10. MVP Acceptance Checklist", level=1)
    add_bullets(doc, [
        "A lead can be created from manual entry and stored with all required fields.",
        "The system flags duplicate phone and property records before creating a new lead.",
        "Clicking Secure Lead sends outreach placeholders or real integration calls, logs activity, and moves the lead to Day 1.",
        "Seller response cancels active contact automation and moves the lead to Contacted.",
        "Intake submission creates structured property data and starts evaluation.",
        "Buy box rules calculate a visible Deal Score and recommended strategy.",
        "Offer status changes are logged and can trigger offer follow-up.",
        "Long-term and forever follow-up create future tasks without deleting the lead.",
        "Dashboard metrics reflect lead, task, offer, contract, and follow-up state.",
        "Role permissions prevent virtual assistants from changing financial fields or offer approvals.",
    ])


def main():
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_section_overview(doc)
    add_domain_model(doc)
    add_states(doc)
    add_workflows(doc)
    add_buy_box(doc)
    add_screens_permissions(doc)
    add_notifications_ai_build(doc)
    doc.save(OUT)


if __name__ == "__main__":
    main()
